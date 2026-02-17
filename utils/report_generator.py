from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import date  # <--- ИСПРАВЛЕН ИМПОРТ
from database.models import Student, EducationalPlan

def generate_word_report(student: Student, plan: EducationalPlan, items: list, logs: list = None) -> BytesIO:
    """
    Генерирует документ Word с индивидуальным планом и журналом.
    """
    doc = Document()

    # --- Настройка стилей ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- Заголовок ---
    head = doc.add_heading(f'ИОМ: {student.full_name}', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1. Общие сведения ---
    doc.add_heading('1. Паспорт программы', level=1)
    p = doc.add_paragraph()
    p.add_run('ФИО ребенка: ').bold = True
    p.add_run(f"{student.full_name}\n")
    p.add_run('Дата рождения: ').bold = True
    p.add_run(f"{student.birth_date.strftime('%d.%m.%Y')}\n")
    p.add_run('Диагноз/Статус: ').bold = True
    p.add_run(f"{student.diagnosis_code}\n")
    
    # Данные о программе
    p.add_run('Цель маршрута: ').bold = True
    p.add_run(f"{plan.goal_description}\n")
    p.add_run('Период реализации: ').bold = True
    # Проверка на случай, если даты не заданы
    s_date = plan.start_date.strftime('%d.%m.%Y') if plan.start_date else "—"
    e_date = plan.end_date.strftime('%d.%m.%Y') if plan.end_date else "—"
    p.add_run(f"с {s_date} по {e_date}\n")

    # --- 2. Таблица упражнений (План) ---
    doc.add_heading('2. Содержание коррекционной работы', level=1)
    
    # Создаем таблицу: №, Навык, Упражнение, Режим
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Заголовки
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Целевой навык'
    hdr_cells[2].text = 'Упражнение/Методика'
    hdr_cells[3].text = 'Режим занятий'

    # Заполняем строками
    for idx, item in enumerate(items):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx + 1)
        skill_name = item.exercise.skill.name if item.exercise.skill else "Общее"
        row_cells[1].text = skill_name
        row_cells[2].text = item.exercise.title
        row_cells[3].text = str(item.frequency) if item.frequency else "По графику"

    # --- 3. Журнал выполнения (Таблица результатов) ---
    if logs:
        doc.add_heading('3. Мониторинг выполнения (Журнал)', level=1)
        doc.add_paragraph("Ниже представлена история занятий и динамика оценок.")
        
        # Создаем таблицу журнала
        log_table = doc.add_table(rows=1, cols=5)
        log_table.style = 'Table Grid'
        
        l_hdr = log_table.rows[0].cells
        l_hdr[0].text = 'Дата'
        l_hdr[1].text = 'Упражнение'
        l_hdr[2].text = 'Статус'
        l_hdr[3].text = 'Балл'
        l_hdr[4].text = 'Заметка'

        # Словарь для красивого статуса
        status_map = {"completed": "Вып.", "failed": "Не спр.", "skipped": "Проп."}

        for log in logs:
            row = log_table.add_row().cells
            row[0].text = log.date.strftime('%d.%m')
            # Получаем название упражнения
            row[1].text = log.item.exercise.title
            # Статус
            status_text = status_map.get(log.status.value, log.status.value)
            row[2].text = status_text
            # Оценка
            row[3].text = str(log.performance_score)
            # Комментарий
            row[4].text = log.teacher_notes if log.teacher_notes else "-"
    else:
        doc.add_heading('3. Мониторинг', level=1)
        doc.add_paragraph("Записи в журнале отсутствуют.")

    # --- Подвал ---
    doc.add_paragraph("\n")
    footer = doc.add_paragraph()
    # ИСПРАВЛЕНИЕ: Используем date.today()
    footer.add_run(f"Дата формирования отчета: {date.today().strftime('%d.%m.%Y')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer